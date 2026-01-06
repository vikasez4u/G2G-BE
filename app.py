import os
from io import BytesIO
from PIL import Image
from docx import Document
from langchain_ollama import ChatOllama
from langchain_chroma import Chroma
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.prompts import PromptTemplate
from langchain.chains.combine_documents import create_stuff_documents_chain
from langchain.chains import create_retrieval_chain
from langchain_community.embeddings.fastembed import FastEmbedEmbeddings
from langchain.docstore.document import Document as LCDocument
from xml.etree.ElementTree import tostring
from functools import lru_cache
import pyodbc

#import streamlit as st

CHROMA_DB_DIR = "./sql_chroma_db"
DOCUMENTS_FOLDER = "./documents"
Collection_Name = "g2g_docs"

embedder = FastEmbedEmbeddings(model_name="BAAI/bge-small-en-v1.5")
#vector_store = Chroma(persist_directory=CHROMA_DB_DIR, embedding_function=embedder,collection_name=Collection_Name)
#print("Vector store loaded:", vector_store._collection.count())

def extract_text_image_link_pairs(doc_path):
    from lxml import etree
    etree.register_namespace("a", "http://schemas.openxmlformats.org/drawingml/2006/main")
    etree.register_namespace("r", "http://schemas.openxmlformats.org/officeDocument/2006/relationships")

    doc = Document(doc_path)
    text_chunks, image_chunks, link_chunks = [], [], []
    rels = doc.part.rels
    images = {}

    for rel in rels.values():
        if "image" in rel.target_ref:
            try:
                image_data = rel.target_part.blob
                img = Image.open(BytesIO(image_data)).convert("RGB")
                images[rel.rId] = img
            except Exception as e:
                print("Image error:", e)

    for para in doc.paragraphs:
        para_text = para.text.strip()
        text_chunks.append(para_text)

        inline_images = []
        if para._element.xpath(".//w:drawing"):
            for drawing in para._element.xpath(".//w:drawing"):
                # embed_id = drawing.xpath(".//a:blip/@r:embed")
                drawing_xml = etree.fromstring(tostring(drawing))
                embed_id = drawing_xml.xpath(
                    ".//a:blip/@r:embed",
                    namespaces={
                        "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
                        "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
                    }
                )

                for rId in embed_id:
                    if rId in images:
                        inline_images.append(images[rId])
        image_chunks.append(inline_images)

        para_links = []
        for hyperlink in para._element.xpath(".//w:hyperlink"):
            rId = hyperlink.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id")
            if rId and rId in rels:
                para_links.append(rels[rId].target_ref)
        link_chunks.append(para_links)

    for i, para in enumerate(doc.paragraphs):
        drawings = para._element.xpath(".//w:drawing")


    print(f"Extracted {len(text_chunks)} paragraphs, {sum(len(imgs) for imgs in image_chunks)} images, and {sum(len(links) for links in link_chunks)} links from {doc_path}")
    return list(zip(text_chunks, image_chunks, link_chunks))

def ingest(): 
    if os.path.exists(CHROMA_DB_DIR):
        # Check if there are any .docx files in the documents folder
        docx_files = [f for f in os.listdir(DOCUMENTS_FOLDER) if f.endswith('.docx')]
        if not docx_files:
            print("Chroma DB already exists, but no documents found.")
        else:
            # Check if all docx files are already present in the DB collection
            db = Chroma(persist_directory=CHROMA_DB_DIR, embedding_function=embedder, collection_name=Collection_Name)
            existing_metadatas = db.get(include=['metadatas'])['metadatas']
            existing_sources = set()
            for meta_list in existing_metadatas:
                if  "source" in meta_list and meta_list["source"]:
                        existing_sources.add(meta_list["source"])
            missing_files = [f for f in docx_files if f not in existing_sources]
            print(f"Missing files: {missing_files}")
            if not missing_files:
                print(f"Chroma DB already exists. All {len(docx_files)} document(s) already ingested: {docx_files}. Skipping ingestion.")
                return
            else:
                print(f"Chroma DB exists, but new documents found: {missing_files}. Proceeding with ingestion.")

    print("Starting ingestion...")

    all_chunks = []
    
    splitter = RecursiveCharacterTextSplitter(chunk_size=800, chunk_overlap=200)

    for filename in os.listdir(DOCUMENTS_FOLDER):
        if filename.endswith(".docx"):
            path = os.path.join(DOCUMENTS_FOLDER, filename)
            para_triplets = extract_text_image_link_pairs(path)
            for i, (para_text, _, _) in enumerate(para_triplets):
                if len(para_text.strip()) < 5:
                    continue
                splits = splitter.split_text(para_text)
                for c in splits:
                    all_chunks.append(LCDocument(
                        page_content=c,
                        metadata={"source": filename, "para_index": i}
                    ))
    print(f"📄 Total chunks being ingested: {len(all_chunks)}")
    db = Chroma.from_documents(documents=all_chunks, embedding=embedder, persist_directory=CHROMA_DB_DIR, collection_name=Collection_Name)
    #db.persist()
    print("✅ Chroma DB created and persisted.")

def build_chain():
    print("Building Chain")
    model = ChatOllama(model="llama3.2", temperature=0.4,
                       options={
            "num_predict": 200,
            "top_p": 0.9,
            "repeat_penalty": 1.1
        }
    )
    prompt = PromptTemplate.from_template(
            f"""
    You are a professional and friendly virtual assistant for Accenture.

    Tone: Warm, formal, consultanting, deep & clear thinker.

    Behavior Guidelines:
    - Greeting or Hi or Hello Message should appears at Start of a new session or chat but not for each message.
    - Base every answer strictly on the provided documents—do not use prior knowledge.
    - If unsure, say "Reach out to Respective POCs"—no assumptions or invented facts.
    - Be concise and contextual.
    - Give detailed explanations only when requested.
    - Use bullet points for multi-part answers.
    - Use emojis sparingly and only when enhancing clarity or warmth.
    - Include links/images from provided documents only if clearly relevant to the user's question.
    - Offer help proactively only if the conversation is just beginning.
    - End with a polite thank-you and positive closing at End of the session or chat but not for each and every message.

    User Question:
    {{input}}

    Context from Documents:
    {{context}}

    Your Answer:
    """
    )

    embedder = FastEmbedEmbeddings(model_name="BAAI/bge-small-en-v1.5")
    vector_store = Chroma(persist_directory=CHROMA_DB_DIR, embedding_function=embedder,collection_name=Collection_Name)
    print("Vector store loaded:", vector_store._collection.count())
    retriever = vector_store.as_retriever(search_type="mmr", search_kwargs={"k": 6, "fetch_k": 12, "lambda_mult": 1.0})
    doc_chain = create_stuff_documents_chain(model, prompt)
    return create_retrieval_chain(retriever, doc_chain), retriever, model

#@lru_cache(maxsize=128)
#def cached_retrieve(user_input: str):
    #return retriever.get_relevant_documents(user_input)
 #   return retriever.invoke(user_input)


ingest()
#chat_chain, chat_retriever,llm = build_chain()

#query = st.text_input("Ask a question:")
#if query:
#    chat_chain, chat_retriever,llm = build_chain()
#    response = chat_chain.invoke({"input": query})
#    st.write("### Answer")
#    st.write(response)

